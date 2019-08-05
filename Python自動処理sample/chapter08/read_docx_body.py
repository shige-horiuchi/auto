import docx

doc = docx.Document('book-proposal1.docx')
print(f'段落総数: {len(doc.paragraphs)}')

for i, paragraph in enumerate(doc.paragraphs):
    print(f'=' * 72)
    print(f'段落 {i} ({paragraph.style.name})')
    print(paragraph.text)
